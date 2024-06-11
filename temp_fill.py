import base64
from streamlit_javascript import st_javascript
import streamlit as st 
from io import BytesIO
from docx import Document

def fill_cont(template_path,output_path,contract_data):
    # Load the template
    doc = Document(template_path)

    # Fill in the contract data
    for key, value in contract_data.items():
        for paragraph in doc.paragraphs:
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)

    # Save the filled contract
    doc.save(output_path)


uploaded_file = st.file_uploader("Upload Your Contract Here", accept_multiple_files=False, type=['pdf'])

def displayPDF(upl_file, ui_width):
    # Read file as bytes:
    bytes_data = upl_file.getvalue()

    # Convert to utf-8
    base64_pdf = base64.b64encode(bytes_data).decode("utf-8")

    # Embed PDF in HTML
    pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width={str(ui_width)} height={str(ui_width*4/3)} type="application/pdf"></iframe>'

    # Display file
    st.markdown(pdf_display, unsafe_allow_html=True)
col1, col2 = st.columns(spec=[2, 1], gap="large")
if uploaded_file:
    data={}
    with col1:
        ui_width = st_javascript("window.innerWidth")
        displayPDF(uploaded_file, ui_width -10)

    with col2:
        st.write("Contract Details:")
        with st.form("employee_contract_form"):
            st.write("Please fill in the following details:")
            
            company_name = st.text_input("Company Name", key="company_name")
            company_address = st.text_input("Company Address", key="company_address")
            representative_title = st.text_input("Representative Title", key="representative_title")
            title = st.selectbox("Title", options=["Mr", "Mme"], key="title")
            employee_name = st.text_input("Employee Name", key="employee_name")
            id_number = st.text_input("ID Number", key="id_number")
            employee_address = st.text_input("Employee Address", key="employee_address")
            job_title = st.text_input("Job Title", key="job_title")
            start_date = st.date_input("Start Date", key="start_date")
            end_date = st.date_input("End Date", key="end_date")
            workplace_address = st.text_input("Workplace Address", key="workplace_address")
            net_salary_amount = st.number_input("Net Salary Amount (DT)", key="net_salary_amount")
            location = st.text_input("Location", key="location")
            
            submit_button = st.form_submit_button("Generate Contract")
            if submit_button:
                data = {
                    "[Company Name]": company_name,
                    "[Company Address]": company_address,
                    "[Representative Title]": representative_title,
                    "[Title]": title,
                    "[Employee Name]": employee_name,
                    "[ID Number]": id_number,
                    "[Employee Address]": employee_address,
                    "[Job Title]": job_title,
                    "[Start Date]": str(start_date),
                    "[End Date]": str(end_date),
                    "[Workplace Address]": workplace_address,
                    "[Net Salary Amount]": str(net_salary_amount),
                    "[Location]": location
                }
        if data:
            doc = Document(r"C:\Users\Bohmid\Desktop\hiil project\beta version\data\cdd_template.docx")

            # Fill in the contract data
            for key, value in data.items():
                for paragraph in doc.paragraphs:
                    if key in paragraph.text:
                        paragraph.text = paragraph.text.replace(key, value)
            doc_io = BytesIO()
            st.write("Document generated successfully 🎉")
            # Save the filled contract
            doc.save(doc_io)
            st.download_button(
                    label="Download Document",
                    data=doc_io.getvalue(), 
                    file_name="cdd_contract.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
